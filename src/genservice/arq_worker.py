import json

import base64
import io
from docx import Document

from arq.connections import logger
from config import redis_settings
from genservice.functionality.feedback_gen_service import feedback_parts_gen_func
from genservice.functionality.genfunctions import extract_requirements, interview_name_extract
from genservice.functionality.transcript_processing_service import categorized_qa_gen_func
from genservice.functionality.utility import parse_requirements, general_requirements


async def process_transcript_task(
    ctx, json_file, requirements, position, filename, timestamp
):
    redis = ctx['redis']

    try:
        transcript_data = json.loads(json_file)
        tech_requirements = parse_requirements(await extract_requirements(requirements))
        interview_name = await interview_name_extract(filename) + timestamp

        categorized_qa = await categorized_qa_gen_func(
            transcript_data=transcript_data,
            tech_requirements=tech_requirements,
            general_requirements=general_requirements,
        )
        processed_data = await feedback_parts_gen_func(
            categorized_qa, position, tech_requirements
        )

        document = Document()
        document.add_heading(f"Feedback for {interview_name}", level=1)
        document.add_paragraph("Generated Feedback:")
        for section, content in processed_data.items():
            document.add_heading(section, level=2)
            document.add_paragraph(content if isinstance(content, str) else json.dumps(content, indent=4))

        file_like = io.BytesIO()
        document.save(file_like)
        file_like.seek(0)

        file_content = base64.b64encode(file_like.getvalue()).decode("utf-8")

        await redis.set(ctx['job_id'], json.dumps({"filename": f"{interview_name}.docx", "file_content": file_content}), ex=3600)
        logger.info(f"Task completed successfully. Result stored for job ID {ctx['job_id']}")

    except Exception as e:
        logger.exception(f"An error occurred while processing the task: {e}")
        await redis.set(ctx['job_id'], json.dumps({"error": str(e)}), ex=3600)


class WorkerSettings:
    functions = [process_transcript_task]
    redis_settings = redis_settings
    max_jobs = 50
    poll_interval = 1
    timeout = 250
    job_timeout = 1000
    log_level = 'INFO'
